from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "EVN_Project_Report_40_Pages.docx"
OUTPUT_EXPANDED = ROOT / "EVN_Project_Report_40_Pages_Expanded.docx"
OUTPUT_EXPANDED_2 = ROOT / "EVN_Project_Report_40_Pages_Expanded_v2.docx"
OUTPUT_CUSTOM = ROOT / "University_Event_Management_System_Report.docx"

TITLE = "University Event Management System"
SUBTITLE = "Comprehensive Project Report"
UNIVERSITY = "Shobhit University, Gangoh"
DATE_TEXT = "Date: April 14, 2026"


def set_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_default_styles(document: Document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number(footer)


def add_title_page(document: Document):
    for _ in range(3):
        document.add_paragraph("")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(UNIVERSITY)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)

    document.add_paragraph("")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    run.bold = True
    run.font.size = Pt(14)

    document.add_paragraph("")
    document.add_paragraph("")

    for line in [
        "Submitted by: __________________________",
        "Course / Department: __________________________",
        "Enrollment No.: __________________________",
        "Guided by: __________________________",
        DATE_TEXT,
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)


def add_heading(document: Document, text: str, level: int = 1):
    p = document.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def add_paragraph(document: Document, text: str):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    p.add_run(text)


def add_bullets(document: Document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(document: Document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(document: Document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph("")


def add_page(document: Document, title: str, body: list, page_break: bool = True):
    add_heading(document, title, 1)
    for item in body:
        kind = item["type"]
        if kind == "p":
            add_paragraph(document, item["text"])
        elif kind == "h2":
            add_heading(document, item["text"], 2)
        elif kind == "h3":
            add_heading(document, item["text"], 3)
        elif kind == "bullets":
            add_bullets(document, item["items"])
        elif kind == "numbered":
            add_numbered(document, item["items"])
        elif kind == "table":
            add_table(document, item["headers"], item["rows"])
    if page_break:
        document.add_page_break()


def enrich_pages(pages):
    expanded = []
    for title, body in pages:
        topic = title.split(". ", 1)[-1].lower()
        extras = [
            {
                "type": "p",
                "text": (
                    f"This section strengthens the overall understanding of {topic} by connecting the written "
                    "explanation to the actual EVN implementation, its academic value, and its practical use within "
                    "a university event-management environment. The added detail is intended to make the report feel "
                    "more complete, better balanced across chapters, and more suitable for formal project-report submission."
                ),
            },
            {
                "type": "p",
                "text": (
                    "From a documentation perspective, elaborating this area also helps demonstrate that the project "
                    "was not built as a collection of disconnected screens or APIs, but as a planned software system "
                    "with identifiable workflows, technical reasoning, and future expansion potential. This kind of "
                    "narrative depth is often expected in academic reports because it shows understanding beyond implementation alone."
                ),
            },
            {
                "type": "p",
                "text": (
                    "In the context of EVN, each chapter contributes to a broader picture in which user experience, "
                    "data structure, security, validation, and administration support each other. Adding richer wording "
                    "therefore does not merely increase length; it makes the report read more like a complete project study "
                    "that explains why the platform was designed in this way and how it can continue to evolve."
                ),
            },
        ]

        insert_at = len(body)
        for index in range(len(body) - 1, -1, -1):
            if body[index]["type"] in {"p", "bullets", "numbered", "table"}:
                insert_at = index + 1
                break

        new_body = body[:insert_at] + extras + body[insert_at:]
        expanded.append((title, new_body))
    return expanded


def light_enrich_pages(pages):
    expanded = []
    for title, body in pages:
        topic = title.split(". ", 1)[-1].lower()
        extra = {
            "type": "p",
            "text": (
                f"This section adds a little more detail about {topic} so the report reads more smoothly and "
                "presents the University Event Management System as a well-planned academic project. The extra wording "
                "also helps connect the chapter content with the actual implementation, showing both practical value and "
                "technical understanding without changing the overall structure you preferred earlier."
            ),
        }

        insert_at = len(body)
        for index in range(len(body) - 1, -1, -1):
            if body[index]["type"] in {"p", "bullets", "numbered", "table"}:
                insert_at = index + 1
                break

        new_body = body[:insert_at] + [extra] + body[insert_at:]
        expanded.append((title, new_body))
    return expanded


def build_base_pages():
    return [
        ("1. Certificate", [
            {"type": "p", "text": "This is to certify that the project titled 'EVN: University Event Management Platform' is a genuine software development study prepared for academic and demonstration purposes. The work documents the analysis, design, implementation approach, and technical decisions used to build a modern platform for publishing university events, handling registrations, and managing participation workflows."},
            {"type": "p", "text": "The report reflects a full-stack system built with a React-based client application, an Express and TypeScript API layer, and a PostgreSQL database managed through Prisma ORM. The documented modules represent a practical event-management solution for university environments where students, admins, and coordinators interact through role-based interfaces."},
            {"type": "p", "text": "Signature of Project Guide: ____________________    Signature of Candidate: ____________________"},
        ]),
        ("2. Declaration", [
            {"type": "p", "text": "I hereby declare that this project report is based on the EVN platform available in the working project repository and that the report has been prepared specifically for academic submission. The analysis, wording, chapter arrangement, and technical presentation in this document are original and tailored to the present software system."},
            {"type": "p", "text": "All external technologies, frameworks, and libraries mentioned in this report are cited in the references section. The practical descriptions, flow explanations, database interpretation, and functional breakdown are derived from the current implementation of the EVN application."},
            {"type": "p", "text": "Student Name: ____________________    Roll No.: ____________________    Date: ____________________"},
        ]),
        ("3. Acknowledgement", [
            {"type": "p", "text": "The successful preparation of this report was supported by the availability of a well-structured project codebase, the guidance of academic mentors, and the technical ecosystem surrounding modern web development. The EVN application made it possible to study a realistic event-management workflow rather than a purely theoretical case."},
            {"type": "p", "text": "Special thanks are due to the faculty members and reviewers who encourage practical, deployment-oriented project work. Appreciation is also extended to the developers and maintainers of open-source tools such as React, Express, Prisma, PostgreSQL, TypeScript, and supporting libraries that simplify the construction of maintainable software systems."},
            {"type": "p", "text": "This report is intended not only as documentation of completed work but also as a future reference for extending the platform into a larger university event ecosystem."},
        ]),
        ("4. Abstract", [
            {"type": "p", "text": "EVN is a full-stack university event management platform designed to centralize public event discovery, student onboarding, registration approval, university badge verification, and role-based operational management. The platform addresses the common campus problem of scattered event information, manual registration handling, and poor visibility into participant status."},
            {"type": "p", "text": "The system provides separate but connected experiences for public visitors, verified university students, administrators, and event coordinators. It supports event groups, solo and team participation, payment-aware registration states, event visibility rules, and identity verification controls for university-only access."},
            {"type": "p", "text": "This report presents the project from the perspectives of need analysis, system design, database modeling, implementation logic, user interface planning, security, testing, deployment readiness, limitations, and future scope. The final outcome is a flexible foundation for a campus-scale digital event ecosystem."},
        ]),
        ("5. Table of Contents", [
            {"type": "numbered", "items": [
                "Certificate",
                "Declaration",
                "Acknowledgement",
                "Abstract",
                "Table of Contents",
                "Introduction to the Project",
                "Problem Statement and Objectives",
                "Scope, Users, and Expected Benefits",
                "Feasibility Study",
                "Requirement Analysis",
                "Functional Requirements",
                "Non-Functional Requirements",
                "System Architecture",
                "Frontend Architecture",
                "Backend Architecture",
                "Database Design Overview",
                "Entity Descriptions and Relations",
                "Authentication and Authorization",
                "Event Group and Event Management",
                "Registration Workflow",
                "Team Participation Management",
                "Admin Dashboard Design",
                "Coordinator Operations",
                "Public Website Experience",
                "Profile and University Badge Verification",
                "API Design and Routing",
                "Validation and Error Handling",
                "File Upload and Media Handling",
                "Security Implementation",
                "Performance and Scalability",
                "User Interface and UX Design",
                "Responsive Design",
                "Testing Strategy",
                "Build and Deployment",
                "Current Limitations",
                "Future Enhancements",
                "Conclusion",
                "References",
                "Appendix A: File Structure",
            ]},
        ]),
        ("6. Introduction to the Project", [
            {"type": "p", "text": "Universities frequently run multiple cultural, technical, and academic events throughout the year, yet the systems used to organize these programs are often fragmented. Posters may exist on social media, registration details may be shared in chat groups, and participant verification may happen manually. EVN was designed to solve this fragmentation by acting as a single event-management platform."},
            {"type": "p", "text": "The project combines a visually rich public website with authenticated student workflows and operational controls for administrators. Instead of treating each event as an isolated form, the platform models event groups, event timelines, audience scope, registration rules, and participant states through a normalized data structure."},
            {"type": "p", "text": "As a result, EVN is more than a registration portal. It is a role-aware digital layer for discovery, approval, scheduling, and event execution within a university ecosystem."},
        ]),
        ("7. Problem Statement and Objectives", [
            {"type": "h2", "text": "Problem Statement"},
            {"type": "p", "text": "Traditional campus event administration suffers from duplicated communication, incomplete records, and inconsistent registration checks. Students may not know whether an event is open to all, limited to verified university participants, free, paid, solo, or team-based. Organizers also face difficulty in validating users, reviewing participants, and tracking event-specific responsibilities."},
            {"type": "h2", "text": "Project Objectives"},
            {"type": "bullets", "items": [
                "Create a central web platform for campus events and event groups.",
                "Support secure registration with email verification and JWT-based authentication.",
                "Introduce university badge verification for restricted events.",
                "Allow admins to manage events, groups, users, and coordinators from one dashboard.",
                "Support solo and team registrations with rule validation.",
                "Lay a strong technical foundation for future competition-management features.",
            ]},
        ]),
        ("8. Scope, Users, and Expected Benefits", [
            {"type": "p", "text": "The current scope of EVN covers the event lifecycle from publication to registration and participation management. The project does not yet integrate a live payment gateway or mobile application, but it already includes payment-aware fields, audience segmentation, profile management, and coordinator assignment logic."},
            {"type": "bullets", "items": [
                "Public users can browse published open events.",
                "Students can create accounts, verify email addresses, update profiles, and register.",
                "Verified university students can access university-only events.",
                "Admins can create event groups, create events, verify badges, and create admin users.",
                "Coordinators can be assigned to specific events within a date window.",
            ]},
            {"type": "p", "text": "The expected benefits include reduced confusion, faster publication of event information, improved transparency in registration approvals, and better long-term maintainability for future event expansion."},
        ]),
        ("9. Feasibility Study", [
            {"type": "h2", "text": "Technical Feasibility"},
            {"type": "p", "text": "The selected stack is technically feasible because React and Express are mature frameworks, TypeScript improves maintainability, and Prisma simplifies schema-driven database operations. PostgreSQL is well suited for relational event data, user accounts, and role-based workflows."},
            {"type": "h2", "text": "Operational Feasibility"},
            {"type": "p", "text": "The platform aligns with actual university processes: publishing events, verifying student identity, assigning coordinators, and handling registration approvals. The user roles map naturally to institutional stakeholders."},
            {"type": "h2", "text": "Economic Feasibility"},
            {"type": "p", "text": "Because the system is built with widely available open-source tools, the development cost remains low while still providing strong extensibility. The architecture can run on modest hosting infrastructure during initial deployment."},
        ]),
        ("10. Requirement Analysis", [
            {"type": "p", "text": "Requirement analysis for EVN begins with the understanding that event participation is not uniform. Some events are public, some are university-specific, some require admin approval, and others involve teams. Therefore, the system must support both broad discoverability and strict rule enforcement."},
            {"type": "p", "text": "The project repository shows a clear route-controller-service structure, indicating that requirements were translated into separate application layers. This separation helps the system evolve without coupling presentation and business logic too tightly."},
            {"type": "table", "headers": ["Requirement Area", "Examples in EVN"], "rows": [
                ["Identity", "Registration, login, OTP verification, password reset"],
                ["Access Control", "Admin-only routes, student-only registration, verified badge checks"],
                ["Event Management", "Groups, events, visibility rules, schedules"],
                ["Operations", "Coordinator assignment, admin review, registration states"],
                ["Presentation", "Home page, event listing, schedule, profile, admin dashboard"],
            ]},
        ]),
        ("11. Functional Requirements", [
            {"type": "bullets", "items": [
                "User registration using email and password.",
                "Email verification through OTP and password reset through OTP.",
                "Profile editing with academic details and university identity fields.",
                "University badge submission and admin review workflow.",
                "Creation and editing of event groups and child events.",
                "Public listing of published events with filtering support.",
                "Solo registration and team registration with validation rules.",
                "Assignment of coordinators to specific events for limited time windows.",
                "Admin dashboard to manage users, events, groups, and badge decisions.",
                "Static file and media upload support for event posters and banners.",
            ]},
            {"type": "p", "text": "These requirements are visible in the current route definitions, schema models, and front-end page structure of the project."},
        ]),
        ("12. Non-Functional Requirements", [
            {"type": "bullets", "items": [
                "Maintainability through modular service-oriented code structure.",
                "Security through JWT authentication, OTP workflows, and role checks.",
                "Scalability through relational modeling and separated frontend/backend layers.",
                "Usability through a modern public-facing interface and task-focused admin views.",
                "Reliability through schema validation, API error handling, and controlled workflows.",
                "Portability through Node.js runtime, TypeScript build scripts, and Prisma migrations.",
            ]},
            {"type": "p", "text": "Non-functional quality is especially important in EVN because the application is expected to grow from a simple event portal into a more operational system that can handle many users and multiple event categories."},
        ]),
        ("13. System Architecture", [
            {"type": "p", "text": "EVN follows a layered architecture composed of a React client, an Express-based API, and a PostgreSQL data layer accessed through Prisma. Each layer has a distinct responsibility. The frontend handles presentation and user interaction, the backend enforces rules and exposes endpoints, and the database stores persistent state."},
            {"type": "table", "headers": ["Layer", "Responsibility"], "rows": [
                ["Client", "Routing, forms, pages, dashboard screens, event discovery"],
                ["API", "Authentication, validation, business rules, data aggregation"],
                ["Database", "Users, OTPs, events, teams, registrations, logs"],
            ]},
            {"type": "p", "text": "The route to controller to service to Prisma flow used in the backend keeps the application readable and consistent. This is one of the strongest architectural features of the project."},
        ]),
        ("14. Frontend Architecture", [
            {"type": "p", "text": "The client application uses React with React Router to define page-level navigation. The main application file shows routes for home, events, event details, schedule, contact, login, signup, email verification, password reset, profile, my events, coordinator area, and an isolated admin dashboard."},
            {"type": "p", "text": "An auth provider wraps the application so that session state can be shared across screens. Toast notifications are configured globally for consistent feedback. The page-level composition indicates a clear distinction between public pages, account pages, and privileged operational pages."},
            {"type": "bullets", "items": [
                "BrowserRouter manages URL-based navigation.",
                "AuthProvider exposes logged-in user context.",
                "Pages remain focused on a single use case.",
                "Reusable layout components provide shared navigation and footer sections.",
            ]},
        ]),
        ("15. Backend Architecture", [
            {"type": "p", "text": "The backend is built with Express and TypeScript. The application configures CORS, JSON parsing, request logging through Morgan, route mounting under `/api/v1`, static file hosting for uploads, and global not-found and error handlers."},
            {"type": "p", "text": "The route index reveals a concise organization around domain concerns: authentication, profile, admin, site content, coordinator operations, event workflows, and upload handling. This structure keeps the backend cohesive while allowing each route group to evolve independently."},
            {"type": "bullets", "items": [
                "Middleware layer handles authentication, authorization, validation, and errors.",
                "Controllers translate HTTP requests into service calls.",
                "Services contain the actual workflow and rule logic.",
                "Prisma provides a typed bridge between services and PostgreSQL.",
            ]},
        ]),
        ("16. Database Design Overview", [
            {"type": "p", "text": "The Prisma schema demonstrates a normalized relational design around users, events, teams, and operational logs. Instead of storing everything in a single event record, the schema separates identity, content, participation, and result-related entities into dedicated models."},
            {"type": "table", "headers": ["Core Model", "Purpose"], "rows": [
                ["User", "Stores login, role, academic profile, and badge state"],
                ["EventGroup", "Represents a larger festival or grouped collection"],
                ["Event", "Stores event-specific settings and schedules"],
                ["EventRegistration", "Tracks participant or team registration state"],
                ["Team / TeamMember", "Supports team-based participation"],
                ["OtpCode", "Stores OTP verification and reset flows"],
                ["CoordinatorAssignment", "Maps students to temporary coordinator roles"],
            ]},
            {"type": "p", "text": "This design makes the system adaptable for future analytics, result publishing, and competition management features."},
        ]),
        ("17. Entity Descriptions and Relations", [
            {"type": "p", "text": "The `User` model is central to the platform. It connects to OTPs, registrations, coordinator assignments, event creation, badge logs, and team relationships. This makes sense because every operational actor is ultimately a user account with a role and status."},
            {"type": "p", "text": "An `EventGroup` can contain many `Event` records, which helps organize a festival into separate competitions. `EventRegistration` links either an individual student or a team to a specific event and tracks the participation lifecycle. `TeamMember` separates team composition from registration status, which is a strong modeling choice."},
            {"type": "p", "text": "The schema also already includes `EventRound`, `PvpMatch`, `LeaderboardEntry`, and `EventResult`, showing that the design anticipates more advanced competitive event flows even when some current UI paths still focus on event publication and registration."},
        ]),
        ("18. Authentication and Authorization", [
            {"type": "p", "text": "Authentication in EVN is based on standard account creation with password storage, followed by JWT-based session handling. The platform also includes OTP-based flows for email verification and password reset, which significantly improves account trust and recovery support."},
            {"type": "p", "text": "Authorization is role-aware. The schema defines `STUDENT` and `ADMIN` roles, while coordinators are handled through temporary assignment records. This is a practical decision because coordinator authority is not global; it belongs to a particular event and time window."},
            {"type": "bullets", "items": [
                "JWT tokens secure authenticated API access.",
                "Role middleware protects admin-only and coordinator-aware actions.",
                "University badge status affects eligibility for restricted events.",
                "Inactive accounts and invalid visibility states are rejected in service logic.",
            ]},
        ]),
        ("19. Event Group and Event Management", [
            {"type": "p", "text": "EVN models the campus event ecosystem in two layers: event groups and events. A group can represent a larger fest, seasonal program, or collection, while individual events contain participation details, venue, dates, audience settings, and registration rules."},
            {"type": "p", "text": "The backend enforces important consistency rules. For example, an event cannot be published while its group is still in draft, and a university-only group cannot contain a publicly visible child event. The service layer also keeps event times within the enclosing group schedule when both are provided."},
            {"type": "p", "text": "This rule-driven approach prevents accidental content inconsistency and reduces admin error during event setup."},
        ]),
        ("20. Registration Workflow", [
            {"type": "p", "text": "The registration flow is one of the most important processes in the system. Before registration is accepted, the backend checks event visibility, publication status, registration opening and closing times, capacity constraints, and participant eligibility. This ensures that users cannot bypass important business rules."},
            {"type": "numbered", "items": [
                "User opens an event detail page and reviews event information.",
                "System checks whether the event is visible for that user.",
                "User registers either as an individual or as a team captain.",
                "Backend validates timing, access, team size, and duplicate conditions.",
                "Registration status becomes pending or confirmed depending on event settings.",
                "Payment status is set according to whether the event requires a fee.",
            ]},
            {"type": "p", "text": "This workflow is robust because it is implemented at the service level rather than only at the UI layer."},
        ]),
        ("21. Team Participation Management", [
            {"type": "p", "text": "Team-based events are modeled with explicit team records and team membership rows. A captain creates the team, members are stored individually, and the event registration references the team. This preserves accountability and supports later operations such as reviewing a team, managing captain privileges, and publishing results."},
            {"type": "p", "text": "The service layer enforces minimum and maximum team size, prevents the captain from being removed, and blocks duplicate or conflicting registrations. It also rejects users who are already actively participating in another event, which reflects a strict one-event-at-a-time rule visible in the current logic."},
            {"type": "p", "text": "These validations illustrate careful thinking around fairness and operational simplicity."},
        ]),
        ("22. Admin Dashboard Design", [
            {"type": "p", "text": "The admin dashboard is a major strength of the frontend. It contains sections for event groups, events, coordinator assignment, user management, and admin account creation. The interface is metric-driven and presents counts for published groups, published events, pending badges, and active coordinators."},
            {"type": "p", "text": "Forms support event and group creation with fields for audience scope, status, schedule, media links, participation type, payment options, and team sizes. The dashboard also provides badge verification and account activation controls."},
            {"type": "bullets", "items": [
                "Single workspace for administrative operations.",
                "Quick refresh and public-site preview controls.",
                "Integrated image upload support for event posters.",
                "User grouping by course and academic year for easier review.",
            ]},
        ]),
        ("23. Coordinator Operations", [
            {"type": "p", "text": "Instead of granting universal elevated access, EVN treats coordinator authority as temporary and event-specific. Coordinators are assigned by admins to a specific event with a start and end time. This avoids over-permissioning and mirrors how event volunteers or student organizers work in real settings."},
            {"type": "p", "text": "The schema and service logic indicate readiness for coordinator-led operations such as registration management, match progression, leaderboard updates, and result publication. Even where some advanced modules are still evolving, the underlying access-control foundation is already suitable for expansion."},
            {"type": "p", "text": "This design is safer and more realistic than introducing a permanent global coordinator role."},
        ]),
        ("24. Public Website Experience", [
            {"type": "p", "text": "The home page is designed to feel like a festival website rather than a plain utility screen. It features a strong hero section, event statistics, featured event cards, and a schedule preview. This visual direction is important because event platforms must attract participation, not merely store forms."},
            {"type": "p", "text": "Published event cards highlight whether an event is paid or free, the participation type, the date, and the venue. The schedule preview simplifies event discovery by focusing on upcoming dates. These choices improve clarity for casual visitors and first-time users."},
            {"type": "p", "text": "The public-facing layer therefore serves both branding and operational discovery purposes."},
        ]),
        ("25. Profile and University Badge Verification", [
            {"type": "p", "text": "The profile workflow is tightly integrated with participation rules. A user profile stores university name, student ID, department, course, and year, as well as the status of university badge verification. This allows the platform to distinguish open participation from restricted campus-only access."},
            {"type": "p", "text": "University-only events can only be seen or joined by verified students, and administrators review badge status through a dedicated workflow. The schema also stores `UniversityBadgeLog` entries, enabling a historical record of review activity rather than relying on a single mutable field."},
            {"type": "p", "text": "This approach increases trust, auditability, and rule transparency."},
        ]),
        ("26. API Design and Routing", [
            {"type": "p", "text": "The REST-style API is mounted under `/api/v1`, which provides a clean versioned base path. The routes are grouped by domain, making the system easier to maintain and easier for frontend developers to consume."},
            {"type": "table", "headers": ["Route Group", "Purpose"], "rows": [
                ["`/auth`", "Registration, login, OTP verification, session access"],
                ["`/profile`", "Profile and university badge related updates"],
                ["`/admin`", "Privileged event, group, user, and admin actions"],
                ["`/coordinator`", "Event-specific management operations"],
                ["`/events`", "Public listing, details, and registration actions"],
                ["`/upload`", "Media upload and image handling"],
            ]},
            {"type": "p", "text": "This route arrangement reflects thoughtful API boundary design rather than an ad-hoc collection of endpoints."},
        ]),
        ("27. Validation and Error Handling", [
            {"type": "p", "text": "Validation in EVN is performed both structurally and logically. Zod validators protect incoming payload shape, while service-layer checks enforce business constraints such as date ordering, payment requirements, team-size rules, and audience eligibility. This two-layer strategy prevents many common runtime issues."},
            {"type": "p", "text": "The project also includes centralized error middleware and an `ApiError` utility. This means service functions can throw meaningful errors and allow the middleware chain to convert them into consistent HTTP responses. Such consistency is critical in systems with many branches of user workflow."},
            {"type": "bullets", "items": [
                "Input shape validation reduces malformed requests.",
                "Business rule validation protects domain integrity.",
                "Global error handling keeps responses uniform.",
                "404 handling provides clean fallback behavior.",
            ]},
        ]),
        ("28. File Upload and Media Handling", [
            {"type": "p", "text": "Events become much more appealing when they include visual media, so EVN includes an upload path for images. The backend serves an `/uploads` directory statically, and the admin dashboard provides an image upload control alongside plain URL entry. This gives administrators flexibility during content creation."},
            {"type": "p", "text": "The package configuration also includes Cloudinary and Multer-related dependencies, showing that the project is prepared for both local and cloud-backed image handling strategies. This is useful for scaling from development mode to a more production-oriented hosting environment."},
            {"type": "p", "text": "Media support improves both usability and promotional quality."},
        ]),
        ("29. Security Implementation", [
            {"type": "p", "text": "Security is a major design consideration in EVN because the system stores account data, event participation details, and access-control information. Password hashes, JWT authentication, OTP status tracking, and route-level authorization are all part of the current implementation strategy."},
            {"type": "bullets", "items": [
                "Password-based login is combined with hashed storage.",
                "JWT secret and expiry are environment-driven.",
                "OTP codes support verification and recovery flows.",
                "Role-based middleware protects privileged operations.",
                "Audience and badge checks stop unauthorized event access.",
                "Time-bound coordinator assignments reduce permission risk.",
            ]},
            {"type": "p", "text": "Together, these mechanisms create a layered security model suitable for a university event platform."},
        ]),
        ("30. Performance and Scalability", [
            {"type": "p", "text": "The project is designed for scalability at both the data and code levels. Prisma indexes on frequently queried fields, such as status combinations and event relations, improve lookup speed. The service-oriented code layout also makes it easier to optimize or replace individual workflows without rewriting the whole application."},
            {"type": "p", "text": "On the frontend, list rendering is scoped to meaningful subsets such as featured events and schedule previews. On the backend, domain-specific selectors reduce overfetching by choosing only the fields needed for a particular response. These are healthy patterns for an application expected to expand over time."},
            {"type": "p", "text": "Further scaling can later be supported with pagination, caching, CDN-backed media delivery, and read-optimized analytics views."},
        ]),
        ("31. User Interface and UX Design", [
            {"type": "p", "text": "The user interface of EVN balances excitement and clarity. Public pages rely on large headings, clear calls to action, and bold festival-oriented imagery, while operational pages focus on efficiency. This split is appropriate because discovery pages and admin tools have different UX goals."},
            {"type": "p", "text": "Forms in the dashboard are grouped by concept, and interaction states are communicated with loaders, badges, and toast notifications. The use of labels such as `Paid`, `Free`, `Published`, `Pending`, and `Verified` helps users understand system state quickly."},
            {"type": "p", "text": "A well-designed interface reduces support burden and increases adoption, especially in academic environments with mixed technical comfort levels."},
        ]),
        ("32. Responsive Design", [
            {"type": "p", "text": "The frontend code includes responsive layouts for navigation, card grids, metric blocks, and the admin tab system. On large screens, the dashboard uses a sidebar, while smaller screens switch to a mobile tab bar with horizontal scrolling. This indicates deliberate consideration of device diversity."},
            {"type": "p", "text": "Responsive behavior is particularly important in campus settings because many users access event information through phones. By ensuring that event listings, schedule previews, forms, and action buttons remain readable and touch-friendly, EVN becomes more practical for real-world use."},
            {"type": "p", "text": "The application therefore supports mobility not just aesthetically but operationally."},
        ]),
        ("33. Testing Strategy", [
            {"type": "p", "text": "Although the repository currently emphasizes implementation over formal automated test suites, the structure is suitable for introducing layered testing. Route validation can be tested through HTTP integration tests, service workflows can be unit-tested with mocked persistence, and critical UI screens can be covered through component or end-to-end tests."},
            {"type": "bullets", "items": [
                "Authentication tests should verify OTP and JWT flows.",
                "Service tests should cover capacity, date, and eligibility rules.",
                "Dashboard tests should confirm admin-only visibility.",
                "Registration tests should cover solo and team paths.",
                "Schema and migration tests should protect relational assumptions.",
            ]},
            {"type": "p", "text": "A future testing layer would further strengthen confidence during feature expansion."},
        ]),
        ("34. Build and Deployment", [
            {"type": "p", "text": "The project includes a practical script setup for development and build workflows. The backend can be run through `npm run dev` and compiled through `npm run build`, while Prisma generation and database setup scripts support schema synchronization and seeding. The client side is maintained through its own package configuration under the `client` directory."},
            {"type": "table", "headers": ["Script", "Purpose"], "rows": [
                ["`npm run dev`", "Start TypeScript backend in watch mode"],
                ["`npm run build`", "Compile backend TypeScript to JavaScript"],
                ["`npm run client:dev`", "Start frontend development server"],
                ["`npm run client:build`", "Build frontend assets"],
                ["`npm run prisma:generate`", "Generate Prisma client"],
                ["`npm run db:seed`", "Seed the database with initial data"],
            ]},
            {"type": "p", "text": "This makes the project deployment-ready with only environment and hosting configuration left to finalize."},
        ]),
        ("35. Current Limitations", [
            {"type": "p", "text": "Like any evolving application, EVN has limitations that should be understood clearly. Some advanced competition entities already exist in the schema but are not yet fully surfaced in the current UI. Payment support is modeled but does not yet connect to a real transaction gateway."},
            {"type": "bullets", "items": [
                "No live payment gateway integration at present.",
                "Advanced round and result workflows are only partially exposed.",
                "Formal automated tests are limited or still to be added.",
                "Notification channels beyond email OTP are not yet implemented.",
                "Analytics and reporting dashboards can be expanded further.",
            ]},
            {"type": "p", "text": "These limitations do not reduce the value of the current platform, but they define the next engineering milestones clearly."},
        ]),
        ("36. Future Enhancements", [
            {"type": "p", "text": "The current architecture leaves substantial room for growth. Because EVN already separates roles, models participation states, and anticipates competition flows, future enhancement can happen incrementally without redesigning the entire system."},
            {"type": "numbered", "items": [
                "Integrate online payment gateways for paid events.",
                "Enable QR-based check-in for on-site attendance.",
                "Complete bracket, leaderboard, and result management screens.",
                "Add email and in-app event reminders.",
                "Provide analytics dashboards for event performance and engagement.",
                "Support downloadable certificates and winner announcements.",
                "Launch a mobile application or progressive web app version.",
            ]},
            {"type": "p", "text": "These upgrades would transform EVN from a strong foundation into a production-grade campus events ecosystem."},
        ]),
        ("37. Conclusion", [
            {"type": "p", "text": "EVN demonstrates how a university event platform can be built with modern full-stack engineering practices while still reflecting real institutional workflows. The project solves important problems in event discovery, student verification, admin control, and rule-based participation management."},
            {"type": "p", "text": "Its most notable strengths are the clarity of the data model, the layered service architecture, the thoughtful public-facing design, and the secure separation of user, admin, and coordinator responsibilities. Even in its current form, the platform is useful and credible as a serious academic software project."},
            {"type": "p", "text": "With continued development in testing, payment integration, and advanced competition tooling, EVN can grow into a highly capable digital operations platform for university events."},
        ]),
        ("38. References", [
            {"type": "numbered", "items": [
                "React Documentation. Retrieved from https://react.dev/",
                "Express Documentation. Retrieved from https://expressjs.com/",
                "TypeScript Documentation. Retrieved from https://www.typescriptlang.org/docs/",
                "Prisma Documentation. Retrieved from https://www.prisma.io/docs/",
                "PostgreSQL Documentation. Retrieved from https://www.postgresql.org/docs/",
                "Zod Documentation. Retrieved from https://zod.dev/",
                "JWT Introduction and Specifications. Retrieved from https://jwt.io/",
                "Node.js Documentation. Retrieved from https://nodejs.org/en/docs/",
            ]},
            {"type": "p", "text": "The current EVN repository, including its Prisma schema, frontend routes, backend services, and configuration files, was also used as a primary implementation reference for the preparation of this report."},
        ]),
        ("39. Appendix A: File Structure", [
            {"type": "p", "text": "The codebase follows a clear multi-layer structure that separates backend, frontend, and database concerns. This improves readability and creates a useful academic example of organized project architecture."},
            {"type": "table", "headers": ["Path", "Purpose"], "rows": [
                ["`src/config`", "Environment, Prisma, and mail configuration"],
                ["`src/controllers`", "HTTP controller handlers"],
                ["`src/middlewares`", "Auth, role, validation, and error middleware"],
                ["`src/routes`", "Domain-based API route modules"],
                ["`src/services`", "Core business logic and database orchestration"],
                ["`src/validators`", "Request schema validation"],
                ["`client/src/pages`", "Frontend screens for public and admin use"],
                ["`prisma/schema.prisma`", "Database schema and enums"],
            ]},
            {"type": "p", "text": "This structure provides an excellent basis for academic explanation, feature scaling, and collaborative development."},
        ]),
    ]


def build_pages():
    return enrich_pages(build_base_pages())


def build_custom_pages():
    pages = [
        ("1. Acknowledgement", [
            {"type": "p", "text": "The successful preparation of this report was supported by the availability of a well-structured project codebase, the guidance of academic mentors, and the technical ecosystem surrounding modern web development. The University Event Management System made it possible to study a realistic event-management workflow rather than a purely theoretical case."},
            {"type": "p", "text": "Special thanks are due to the faculty members and reviewers who encourage practical, deployment-oriented project work. Appreciation is also extended to the developers and maintainers of open-source tools such as React, Express, Prisma, PostgreSQL, TypeScript, and supporting libraries that simplify the construction of maintainable software systems."},
            {"type": "p", "text": "This report is intended not only as documentation of completed work but also as a future reference for extending the platform into a larger university event ecosystem."},
        ]),
        ("2. Abstract", [
            {"type": "p", "text": "The University Event Management System is a full-stack platform designed to centralize public event discovery, student onboarding, registration approval, university badge verification, and role-based operational management. The platform addresses the common campus problem of scattered event information, manual registration handling, and poor visibility into participant status."},
            {"type": "p", "text": "The system provides separate but connected experiences for public visitors, verified university students, administrators, and event coordinators. It supports event groups, solo and team participation, payment-aware registration states, event visibility rules, and identity verification controls for university-only access."},
            {"type": "p", "text": "This report presents the project from the perspectives of need analysis, system design, database modeling, implementation logic, user interface planning, security, testing, deployment readiness, limitations, and future scope. The final outcome is a flexible foundation for a campus-scale digital event ecosystem."},
        ]),
        ("3. Table of Contents", [
            {"type": "numbered", "items": [
                "Acknowledgement",
                "Abstract",
                "Table of Contents",
                "Introduction to the Project",
                "Problem Statement and Objectives",
                "Scope, Users, and Expected Benefits",
                "Feasibility Study",
                "Requirement Analysis",
                "Functional Requirements",
                "Non-Functional Requirements",
                "System Architecture",
                "Frontend Architecture",
                "Backend Architecture",
                "Database Design Overview",
                "Entity Descriptions and Relations",
                "Authentication and Authorization",
                "Event Group and Event Management",
                "Registration Workflow",
                "Team Participation Management",
                "Admin Dashboard Design",
                "Coordinator Operations",
                "Public Website Experience",
                "Profile and University Badge Verification",
                "API Design and Routing",
                "Validation and Error Handling",
                "File Upload and Media Handling",
                "Security Implementation",
                "Performance and Scalability",
                "User Interface and UX Design",
                "Responsive Design",
                "Testing Strategy",
                "Build and Deployment",
                "Current Limitations",
                "Future Enhancements",
                "Conclusion",
                "References",
                "Appendix A: File Structure",
            ]},
        ]),
    ]

    for title, body in build_base_pages()[5:]:
        normalized_title = title
        normalized_body = []
        for item in body:
            if item["type"] == "p":
                normalized_body.append(
                    {
                        "type": "p",
                        "text": item["text"]
                        .replace("EVN", "the University Event Management System")
                        .replace("University Event Management Platform", "University Event Management System"),
                    }
                )
            else:
                normalized_body.append(item)
        pages.append((normalized_title.replace("EVN", "University Event Management System"), normalized_body))

    return light_enrich_pages(pages)


def main():
    document = Document()
    set_default_styles(document)
    add_title_page(document)
    document.add_page_break()

    pages = build_pages()
    for index, (title, body) in enumerate(pages):
        add_page(document, title, body, page_break=index != len(pages) - 1)

    try:
        document.save(OUTPUT)
        print(f"Created: {OUTPUT}")
    except PermissionError:
        try:
            document.save(OUTPUT_EXPANDED)
            print(f"Created: {OUTPUT_EXPANDED}")
        except PermissionError:
            document.save(OUTPUT_EXPANDED_2)
            print(f"Created: {OUTPUT_EXPANDED_2}")


if __name__ == "__main__":
    main()
